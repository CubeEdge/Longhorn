#!/usr/bin/env python3
"""
DOCX转Markdown脚本 (增强版)
- 保留标题层级
- 转换表格为Markdown格式（完整支持）
- 提取图片并转WebP
- 保留粗体/斜体/列表格式
"""

import sys
import os
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
from PIL import Image
import io
import hashlib
import re

def table_to_markdown(table):
    """转换DOCX表格为Markdown"""
    if not table.rows:
        return ''
    
    md_lines = []
    rows_data = []
    
    # 提取所有行
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            text = cell.text.strip().replace('\n', ' ')
            row_data.append(text)
        rows_data.append(row_data)
    
    if len(rows_data) < 1:
        return ''
    
    # 确保列数一致
    max_cols = max(len(row) for row in rows_data)
    for row in rows_data:
        while len(row) < max_cols:
            row.append('')
    
    # 表头
    header = rows_data[0]
    md_lines.append('| ' + ' | '.join(header) + ' |')
    
    # 分隔线
    md_lines.append('| ' + ' | '.join(['---'] * len(header)) + ' |')
    
    # 数据行
    for row in rows_data[1:]:
        md_lines.append('| ' + ' | '.join(row) + ' |')
    
    return '\n'.join(md_lines)

def process_run_formatting(run):
    """处理文本格式"""
    text = run.text
    if not text:
        return ''
    
    # 粗体
    if run.bold:
        text = f'**{text}**'
    
    # 斜体
    if run.italic:
        text = f'*{text}*'
    
    return text

def extract_images_from_docx(docx_path, images_dir):
    """提取DOCX中的图片"""
    doc = Document(docx_path)
    Path(images_dir).mkdir(parents=True, exist_ok=True)
    
    image_map = {}  # rId -> filepath
    
    # 遍历所有图片关系
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_bytes = rel.target_part.blob
                img_hash = hashlib.md5(image_bytes).hexdigest()[:12]
                
                # 转WebP
                img = Image.open(io.BytesIO(image_bytes))
                
                # 处理颜色模式
                if img.mode in ('RGBA', 'LA', 'P'):
                    bg = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P': 
                        img = img.convert('RGBA')
                    if img.mode in ('RGBA', 'LA'): 
                        bg.paste(img, mask=img.split()[-1])
                    img = bg
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                filename = f"img_{img_hash}.webp"
                filepath = os.path.join(images_dir, filename)
                img.save(filepath, "WEBP", quality=85, method=6)
                
                image_map[rel.rId] = f"/data/knowledge_images/{filename}"
                print(f"      提取图片: {filename}")
            except Exception as e:
                print(f"      图片处理失败: {e}")
    
    return image_map

def convert_docx_to_markdown(docx_path, output_md_path, images_dir):
    """转换DOCX为Markdown"""
    print(f"[1/4] 读取DOCX文件: {docx_path}")
    doc = Document(docx_path)
    
    print(f"[2/4] 提取图片...")
    image_map = extract_images_from_docx(docx_path, images_dir)
    
    print(f"\n[3/4] 转换内容...")
    markdown_lines = []
    
    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith('p'):
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if not para:
                continue
            
            # 检测标题
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                title = para.text.strip()
                if title:
                    markdown_lines.append(f"{'#' * level} {title}\n")
            else:
                # 普通段落
                para_text = ''
                for run in para.runs:
                    para_text += process_run_formatting(run)
                
                # 检查图片
                for drawing in para._element.findall('.//' + qn('w:drawing')):
                    for blip in drawing.findall('.//' + qn('a:blip')):
                        embed_id = blip.get(qn('r:embed'))
                        if embed_id in image_map:
                            para_text += f"\n\n![Image]({image_map[embed_id]})\n\n"
                
                if para_text.strip():
                    markdown_lines.append(para_text.strip() + '\n\n')
        
        # 处理表格
        elif element.tag.endswith('tbl'):
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            
            if table:
                md_table = table_to_markdown(table)
                if md_table:
                    markdown_lines.append(md_table + '\n\n')
    
    # 合并内容
    markdown_content = ''.join(markdown_lines)
    
    # 清理多余空行
    markdown_content = re.sub(r'\n{4,}', '\n\n\n', markdown_content)
    
    print(f"[4/4] 保存Markdown文件: {output_md_path}")
    with open(output_md_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    # 统计
    lines = markdown_content.split('\n')
    heading_count = len([l for l in lines if l.strip().startswith('#')])
    image_count = len([l for l in lines if '![' in l])
    table_lines = [l for l in lines if l.strip().startswith('|')]
    table_count = len([l for l in table_lines if '---' in l])  # 数表格，不是行
    
    print(f"\n✅ 转换完成！")
    print(f"📊 统计:")
    print(f"   - 总行数: {len(lines)}")
    print(f"   - 标题数: {heading_count}")
    print(f"   - 图片数: {image_count}")
    print(f"   - 表格数: {table_count}")
    
    return markdown_content

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python3 docx_to_markdown.py <docx_path> <output_md_path> <images_dir>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_md_path = sys.argv[2]
    images_dir = sys.argv[3]
    
    if not os.path.exists(docx_path):
        print(f"❌ 文件不存在: {docx_path}")
        sys.exit(1)
    
    convert_docx_to_markdown(docx_path, output_md_path, images_dir)
