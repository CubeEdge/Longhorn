#!/usr/bin/env python3
"""
DOCX转HTML脚本
- 保留标题层级 (h1~h6)
- 转换表格为HTML table
- 提取图片并转WebP
- 保留粗体/斜体/列表格式
"""

import sys
import os
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
from PIL import Image
import io
import hashlib
import re
from html import escape

def table_to_html(table):
    """转换DOCX表格为HTML table，保留格式"""
    if not table.rows:
        return ''
    
    html_lines = ['<table class="wiki-table">']
    
    for row_idx, row in enumerate(table.rows):
        html_lines.append('  <tr>')
        for cell in row.cells:
            # 处理单元格内的段落和格式
            cell_content = ''
            for para in cell.paragraphs:
                para_text = ''
                for run in para.runs:
                    text = escape(run.text)
                    if not text:
                        continue
                    # 保留粗体、斜体格式
                    if run.bold:
                        text = f'<strong>{text}</strong>'
                    if run.italic:
                        text = f'<em>{text}</em>'
                    para_text += text
                if para_text.strip():
                    cell_content += para_text + '<br>'
            
            cell_content = cell_content.rstrip('<br>')
            
            # 第一行作为表头
            if row_idx == 0:
                html_lines.append(f'    <th>{cell_content}</th>')
            else:
                html_lines.append(f'    <td>{cell_content}</td>')
        html_lines.append('  </tr>')
    
    html_lines.append('</table>')
    return '\n'.join(html_lines)

def process_run_formatting(run):
    """处理文本格式，输出HTML"""
    text = run.text
    if not text:
        return ''
    
    # 转义HTML特殊字符
    text = escape(text)
    
    # 粗体
    if run.bold:
        text = f'<strong>{text}</strong>'
    
    # 斜体
    if run.italic:
        text = f'<em>{text}</em>'
    
    return text

def extract_images_from_docx(docx_path, images_dir):
    """提取DOCX中的图片"""
    doc = Document(docx_path)
    Path(images_dir).mkdir(parents=True, exist_ok=True)
    
    image_map = {}  # rId -> filepath
    
    # 遍历所有图片关系
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_bytes = rel.target_part.blob
                img_hash = hashlib.md5(image_bytes).hexdigest()[:12]
                
                # 转WebP
                img = Image.open(io.BytesIO(image_bytes))
                
                # 处理颜色模式
                if img.mode in ('RGBA', 'LA', 'P'):
                    bg = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P': 
                        img = img.convert('RGBA')
                    if img.mode in ('RGBA', 'LA'): 
                        bg.paste(img, mask=img.split()[-1])
                    img = bg
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                filename = f"img_{img_hash}.webp"
                filepath = os.path.join(images_dir, filename)
                img.save(filepath, "WEBP", quality=85, method=6)
                
                image_map[rel.rId] = f"/data/knowledge_images/{filename}"
                print(f"      提取图片: {filename}")
            except Exception as e:
                print(f"      图片处理失败: {e}")
    
    return image_map

def convert_docx_to_html(docx_path, output_html_path, images_dir):
    """转换DOCX为HTML"""
    print(f"[1/4] 读取DOCX文件: {docx_path}")
    doc = Document(docx_path)
    
    print(f"[2/4] 提取图片...")
    image_map = extract_images_from_docx(docx_path, images_dir)
    
    print(f"\n[3/4] 转换内容...")
    html_lines = []
    
    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith('p'):
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break
            
            if not para:
                continue
            
            # 检测标题
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                title = escape(para.text.strip())
                if title:
                    html_lines.append(f"<h{level}>{title}</h{level}>\n")
            else:
                # 普通段落
                para_text = ''
                for run in para.runs:
                    para_text += process_run_formatting(run)
                
                # 检查图片
                images_html = []
                for drawing in para._element.findall('.//' + qn('w:drawing')):
                    for blip in drawing.findall('.//' + qn('a:blip')):
                        embed_id = blip.get(qn('r:embed'))
                        if embed_id in image_map:
                            images_html.append(f'<img src="{image_map[embed_id]}" alt="Image" />')
                
                if para_text.strip():
                    html_lines.append(f"<p>{para_text.strip()}</p>\n")
                
                for img_html in images_html:
                    html_lines.append(f"<p>{img_html}</p>\n")
        
        # 处理表格
        elif element.tag.endswith('tbl'):
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            
            if table:
                html_table = table_to_html(table)
                if html_table:
                    html_lines.append(html_table + '\n\n')
    
    # 合并内容
    html_content = ''.join(html_lines)
    
    # 清理多余空行
    html_content = re.sub(r'\n{3,}', '\n\n', html_content)
    
    print(f"[4/4] 保存HTML文件: {output_html_path}")
    with open(output_html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    # 统计
    heading_count = len(re.findall(r'<h[1-6]>', html_content))
    image_count = len(re.findall(r'<img ', html_content))
    table_count = len(re.findall(r'<table', html_content))
    
    print(f"\n✅ 转换完成！")
    print(f"📊 统计:")
    print(f"   - 总字符数: {len(html_content)}")
    print(f"   - 标题数: {heading_count}")
    print(f"   - 图片数: {image_count}")
    print(f"   - 表格数: {table_count}")
    
    return html_content

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python3 docx_to_html.py <docx_path> <output_html_path> <images_dir>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_html_path = sys.argv[2]
    images_dir = sys.argv[3]
    
    if not os.path.exists(docx_path):
        print(f"❌ 文件不存在: {docx_path}")
        sys.exit(1)
    
    convert_docx_to_html(docx_path, output_html_path, images_dir)
