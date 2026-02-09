#!/usr/bin/env python3
"""
从Word文档导入MAVO Edge 6K操作说明书到知识库
- 提取文字和结构
- 提取图片并保存
- 准确匹配图片到文字位置
"""

import sys
import os
import sqlite3
import hashlib
import re
from pathlib import Path
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
import io

# 配置
DOCX_PATH = "/Users/Kine/Documents/Kinefinity/KineCore/Pool/qoder/Longhorn/input docs/MAVO Edge 6K操作说明书(KineOS8.0)_C34-102-8016_2024.12.19_v0.1_Jiulong.docx"
DB_PATH = "/Users/Kine/Documents/Kinefinity/KineCore/Pool/qoder/Longhorn/server/longhorn.db"
IMAGE_OUTPUT_DIR = "/Users/Kine/Documents/Kinefinity/KineCore/Pool/qoder/Longhorn/server/data/knowledge_images"

Path(IMAGE_OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

def extract_images_from_docx(doc):
    """从Word文档提取所有图片"""
    images = []
    image_parts = {}
    
    # 获取所有图片关系
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_parts[rel.rId] = rel.target_part
    
    print(f"找到 {len(image_parts)} 个图片引用")
    
    # 提取图片
    for rid, image_part in image_parts.items():
        try:
            image_bytes = image_part.blob
            img = Image.open(io.BytesIO(image_bytes))
            width, height = img.size
            
            # 跳过小图标
            if width < 50 or height < 50:
                continue
            
            # 生成文件名
            img_hash = hashlib.md5(image_bytes).hexdigest()[:12]
            filename = f"edge6k_docx_{img_hash}.png"
            filepath = os.path.join(IMAGE_OUTPUT_DIR, filename)
            
            # 保存为PNG
            if img.mode not in ('RGB', 'RGBA', 'L'):
                img = img.convert('RGB')
            img.save(filepath, "PNG")
            
            images.append({
                'rid': rid,
                'filename': filename,
                'width': width,
                'height': height,
                'path': f'/data/knowledge_images/{filename}'
            })
            
            print(f"  ✓ {filename} ({width}x{height})")
            
        except Exception as e:
            print(f"  ⚠ 图片提取失败 ({rid}): {e}")
            continue
    
    return images

def get_paragraph_images(paragraph):
    """获取段落中的图片引用ID"""
    image_rids = []
    
    # 检查段落的XML
    for run in paragraph.runs:
        # 查找图片引用
        for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            for blip in drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed:
                    image_rids.append(embed)
    
    return image_rids

def is_heading(paragraph):
    """判断段落是否为标题"""
    if paragraph.style.name.startswith('Heading'):
        return True
    
    # 检查是否有章节编号格式
    text = paragraph.text.strip()
    if re.match(r'^\d+(?:\.\d+)*\s+.+', text):
        # 检查字体大小或加粗
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold or (first_run.font.size and first_run.font.size.pt > 12):
                return True
    
    return False

def parse_chapters_from_docx(doc, image_map):
    """从Word文档提取章节内容"""
    chapters = []
    current_chapter = None
    
    # 章节标题模式
    chapter_pattern = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)$')
    
    for element in doc.element.body:
        # 处理段落
        if isinstance(element, CT_P):
            paragraph = Paragraph(element, doc)
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # 检查是否为章节标题
            if is_heading(paragraph):
                match = chapter_pattern.match(text)
                if match:
                    # 保存上一章节
                    if current_chapter and current_chapter['content']:
                        chapters.append(current_chapter)
                    
                    # 开始新章节
                    chapter_num = match.group(1)
                    chapter_title = match.group(2)
                    current_chapter = {
                        'number': chapter_num,
                        'title': f"MAVO Edge 6K: {chapter_title}",
                        'full_title': text,
                        'content': "",
                        'images': []
                    }
                    continue
            
            # 添加到当前章节
            if current_chapter:
                # 检查段落中的图片
                img_rids = get_paragraph_images(paragraph)
                
                # 添加文字
                current_chapter['content'] += text + "\n\n"
                
                # 添加图片引用
                for rid in img_rids:
                    if rid in image_map:
                        img_info = image_map[rid]
                        current_chapter['content'] += f"![{img_info['filename']}]({img_info['path']})\n\n"
                        current_chapter['images'].append(img_info)
        
        # 处理表格（暂时跳过，可以后续增强）
        elif isinstance(element, CT_Tbl):
            pass
    
    # 保存最后一章
    if current_chapter and current_chapter['content']:
        chapters.append(current_chapter)
    
    return chapters

def main():
    print("=" * 70)
    print("从Word文档导入 MAVO Edge 6K 操作说明书")
    print("=" * 70)
    
    # 1. 打开Word文档
    print(f"\n📄 打开Word文档...")
    doc = Document(DOCX_PATH)
    print(f"   段落数: {len(doc.paragraphs)}")
    print(f"   表格数: {len(doc.tables)}")
    
    # 2. 提取图片
    print(f"\n📸 提取图片...")
    images = extract_images_from_docx(doc)
    print(f"   共提取 {len(images)} 张有效图片")
    
    # 创建图片映射
    image_map = {img['rid']: img for img in images}
    
    # 3. 提取章节
    print(f"\n📖 提取章节内容...")
    chapters = parse_chapters_from_docx(doc, image_map)
    print(f"   共提取 {len(chapters)} 个章节")
    
    # 显示章节预览
    for i, ch in enumerate(chapters[:5]):
        print(f"   {i+1}. {ch['number']} {ch['title'][:40]}... ({len(ch['images'])}张图)")
    if len(chapters) > 5:
        print(f"   ... 还有 {len(chapters)-5} 个章节")
    
    # 4. 连接数据库
    print(f"\n💾 连接数据库...")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # 5. 清空旧数据（已经清空过，这里确认）
    cursor.execute("DELETE FROM knowledge_articles WHERE category = 'Manual'")
    print(f"   已清空Manual分类")
    
    # 6. 获取admin用户ID
    admin_user = cursor.execute("SELECT id FROM users WHERE username = 'admin' LIMIT 1").fetchone()
    admin_id = admin_user[0] if admin_user else 1
    
    # 7. 插入新文章
    print(f"\n✍️  插入文章...")
    insert_sql = """
        INSERT INTO knowledge_articles (
            title, slug, summary, content,
            category, subcategory, product_line, product_models,
            visibility, status, published_at, created_by, created_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, datetime('now'), ?, datetime('now'))
    """
    
    for i, chapter in enumerate(chapters):
        # 生成slug
        slug = f"edge-6k-{chapter['number'].replace('.', '-')}"
        
        # 生成摘要（取前200字符）
        summary = chapter['content'][:200].replace('\n', ' ').strip()
        
        try:
            cursor.execute(insert_sql, (
                chapter['title'],
                slug,
                summary,
                chapter['content'],
                'Manual',
                '操作手册',
                'MAVO Edge',
                'MAVO Edge 6K',
                'Public',
                'Published',
                admin_id
            ))
            print(f"   ✓ {chapter['number']} {chapter['title'][:50]}")
        except sqlite3.IntegrityError as e:
            # 如果slug重复，添加序号
            slug = f"{slug}-{i+1:03d}"
            cursor.execute(insert_sql, (
                chapter['title'],
                slug,
                summary,
                chapter['content'],
                'Manual',
                '操作手册',
                'MAVO Edge',
                'MAVO Edge 6K',
                'Public',
                'Published',
                admin_id
            ))
            print(f"   ✓ {chapter['number']} {chapter['title'][:50]} (slug: {slug})")
    
    conn.commit()
    
    # 8. 验证
    count = cursor.execute("SELECT COUNT(*) FROM knowledge_articles WHERE category = 'Manual'").fetchone()[0]
    print(f"\n✅ 完成！共导入 {count} 篇文章")
    
    # 显示示例
    test = cursor.execute("""
        SELECT title, SUBSTR(content, 1, 200)
        FROM knowledge_articles
        WHERE category = 'Manual'
        ORDER BY id
        LIMIT 1
    """).fetchone()
    
    if test:
        print(f"\n📝 示例文章:")
        print(f"   标题: {test[0]}")
        print(f"   内容: {test[1][:100]}...")
    
    conn.close()
    print("\n" + "=" * 70)

if __name__ == '__main__':
    main()
